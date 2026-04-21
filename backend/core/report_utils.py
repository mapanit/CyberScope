#!/usr/bin/env python3
"""
Утилиты для создания унифицированных отчетов для всех инструментов
Каждый инструмент сохраняет отчеты в своей папке /reports
Форматы: TXT и JSON
"""

import json
import os
from pathlib import Path
from datetime import datetime, timedelta
from typing import Dict, Any, Optional, List, Tuple
from abc import ABC, abstractmethod
try:
    import nvdlib
    NVDLIB_AVAILABLE = True
except ImportError:
    NVDLIB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ReportBase(ABC):
    """Базовый класс для создания отчетов"""

    def __init__(self, tool_name: str, target: str, reports_base_dir: Optional[Path] = None):
        """
        Инициализация отчета

        Args:
            tool_name: Имя инструмента (e.g., 'wappalyzer', 'scanner', 'nuclei')
            target: Целевой URL или домен
            reports_base_dir: Базовая директория для отчетов
        """
        self.tool_name = tool_name
        self.target = target
        self.scan_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.scan_datetime = datetime.now().isoformat()

        # Настройка директорий
        if reports_base_dir is None:
            reports_base_dir = Path(__file__).parent.parent / "reports"
        else:
            reports_base_dir = Path(reports_base_dir)

        # Создаем папку для инструмента
        self.tool_dir = reports_base_dir / tool_name
        self.json_dir = self.tool_dir / "json"
        self.txt_dir = self.tool_dir / "txt"

        # Создаем директории
        for directory in [self.json_dir, self.txt_dir]:
            directory.mkdir(parents=True, exist_ok=True)
            # Установка прав доступа для директорий
            os.chmod(directory, 0o755)

        self.filename_base = f"{tool_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    @abstractmethod
    def get_json_report(self) -> Dict[str, Any]:
        """Получить отчет в формате JSON"""
        pass

    def save_json_report(self, data: Dict[str, Any]) -> str:
        """Сохранить JSON отчет"""
        report_path = self.json_dir / f"{self.filename_base}.json"

        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        # Установка прав доступа для чтения всем пользователям
        os.chmod(report_path, 0o644)

        print(f"[+] JSON отчет {self.tool_name} сохранен: {report_path}")
        return str(report_path)

    def save_txt_report(self, content: str) -> str:
        """Сохранить TXT отчет"""
        report_path = self.txt_dir / f"{self.filename_base}.txt"

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Установка прав доступа для чтения всем пользователям
        os.chmod(report_path, 0o644)

        print(f"[+] TXT отчет {self.tool_name} сохранен: {report_path}")
        return str(report_path)


class CombinedReport:
    """Класс для объединения отчетов от всех инструментов в JSON и TXT формате"""

    def __init__(self, scan_id: str, reports_base_dir: Optional[Path] = None):
        """
        Инициализация объединенного отчета

        Args:
            scan_id: ID сканирования
            reports_base_dir: Базовая директория для отчетов
        """
        self.scan_id = scan_id
        self.scan_datetime = datetime.now().isoformat()

        if reports_base_dir is None:
            reports_base_dir = Path(__file__).parent.parent / "reports"
        else:
            reports_base_dir = Path(reports_base_dir)

        self.combined_dir = reports_base_dir / "combined"
        self.json_dir = self.combined_dir / "json"
        self.txt_dir = self.combined_dir / "txt"
        self.word_dir = self.combined_dir / "word"

        # Создаем директории
        for directory in [self.json_dir, self.txt_dir, self.word_dir]:
            directory.mkdir(parents=True, exist_ok=True)
            # Установка прав доступа для директорий
            os.chmod(directory, 0o755)

        self.filename_base = f"combined_report_{scan_id}"
        self.tool_reports = {}
        self.targets = set()
        self.tool_files = {}  # Хранит пути к файлам каждого инструмента
        self.scan_start_time = None  # Время начала текущего сканирования

    def set_scan_start_time(self, start_time: datetime):
        """
        Установить время начала сканирования для фильтрации отчетов

        Args:
            start_time: Время начала сканирования
        """
        self.scan_start_time = start_time

    def collect_files_by_time_window(self, reports_base_dir: Optional[Path] = None,
                                     start_time: Optional[datetime] = None,
                                     end_time: Optional[datetime] = None) -> Dict[str, Path]:
        """
        Собрать пути к TXT файлам, созданным в указанном временном окне

        Args:
            reports_base_dir: Базовая директория отчетов
            start_time: Начало временного окна (если None, используется время начала сканирования)
            end_time: Конец временного окна (если None, используется текущее время)

        Returns:
            Dict с путями к TXT файлам для каждого инструмента
        """
        if reports_base_dir is None:
            reports_base_dir = Path(__file__).parent.parent / "reports"
        else:
            reports_base_dir = Path(reports_base_dir)

        # Устанавливаем временные границы
        if start_time is None:
            start_time = self.scan_start_time if self.scan_start_time else datetime.now() - \
                timedelta(minutes=10)

        if end_time is None:
            end_time = datetime.now()

        # Список инструментов для сбора
        tools = ['wappalyzer', 'scanner', 'nuclei',
                 'whois', 'web', 'osint', 'amass', 'katana', 'retire', 'cors', 'ssl-tls', 'dns', 'nmap']

        collected_files = {}

        print(
            f"[*] Поиск отчетов созданных между {start_time.strftime('%H:%M:%S')} и {end_time.strftime('%H:%M:%S')}")

        for tool_name in tools:
            txt_dir = reports_base_dir / tool_name / "txt"
            if txt_dir.exists() and txt_dir.is_dir():
                # Получаем все TXT файлы
                txt_files = list(txt_dir.glob("*.txt"))
                if txt_files:
                    # Фильтруем по времени создания
                    files_in_window = []
                    for f in txt_files:
                        file_time = datetime.fromtimestamp(f.stat().st_mtime)
                        if start_time <= file_time <= end_time:
                            files_in_window.append((f, file_time))

                    if files_in_window:
                        # Берем самый свежий файл в окне
                        latest_file = max(
                            files_in_window, key=lambda x: x[1])[0]
                        collected_files[tool_name] = latest_file
                        print(f"[+] Найден отчет {tool_name}: {latest_file.name} "
                              f"(создан: {datetime.fromtimestamp(latest_file.stat().st_mtime).strftime('%H:%M:%S')})")
                    else:
                        # Если в окне нет файлов, ищем самый свежий файл за последние 5 минут
                        recent_files = [
                            f for f in txt_files
                            if datetime.fromtimestamp(f.stat().st_mtime) >= datetime.now() - timedelta(minutes=5)
                        ]
                        if recent_files:
                            latest_file = max(
                                recent_files, key=lambda p: p.stat().st_mtime)
                            collected_files[tool_name] = latest_file
                            print(
                                f"[*] Найден недавний отчет {tool_name}: {latest_file.name}")

        self.tool_files = collected_files
        return collected_files

    def collect_recent_files(self, reports_base_dir: Optional[Path] = None,
                             recent_minutes: int = 5) -> Dict[str, Path]:
        """
        Собрать пути к недавно созданным TXT файлам

        Args:
            reports_base_dir: Базовая директория отчетов
            recent_minutes: Количество минут для поиска недавних отчетов

        Returns:
            Dict с путями к TXT файлам для каждого инструмента
        """
        if reports_base_dir is None:
            reports_base_dir = Path(__file__).parent.parent / "reports"
        else:
            reports_base_dir = Path(reports_base_dir)

        # Время для фильтрации
        cutoff_time = datetime.now() - timedelta(minutes=recent_minutes)

        # Список инструментов для сбора
        tools = ['wappalyzer', 'scanner', 'nuclei',
                 'whois', 'web', 'osint', 'amass', 'katana', 'retire', 'cors', 'ssl-tls', 'dns', 'nmap']

        collected_files = {}

        print(
            f"[*] Поиск отчетов созданных за последние {recent_minutes} минут (после {cutoff_time.strftime('%H:%M:%S')})")

        for tool_name in tools:
            txt_dir = reports_base_dir / tool_name / "txt"
            if txt_dir.exists() and txt_dir.is_dir():
                # Получаем все TXT файлы
                txt_files = list(txt_dir.glob("*.txt"))
                if txt_files:
                    # Фильтруем по времени (недавние)
                    recent_files = [
                        f for f in txt_files
                        if datetime.fromtimestamp(f.stat().st_mtime) >= cutoff_time
                    ]

                    if recent_files:
                        # Берем самый свежий файл
                        latest_file = max(
                            recent_files, key=lambda p: p.stat().st_mtime)
                        collected_files[tool_name] = latest_file
                        print(f"[+] Найден отчет {tool_name}: {latest_file.name} "
                              f"(создан: {datetime.fromtimestamp(latest_file.stat().st_mtime).strftime('%H:%M:%S')})")
                    else:
                        # Проверяем файлы за последние 10 минут
                        extended_cutoff = datetime.now() - timedelta(minutes=10)
                        extended_files = [
                            f for f in txt_files
                            if datetime.fromtimestamp(f.stat().st_mtime) >= extended_cutoff
                        ]
                        if extended_files:
                            latest_file = max(
                                extended_files, key=lambda p: p.stat().st_mtime)
                            collected_files[tool_name] = latest_file
                            print(
                                f"[*] Найден отчет {tool_name} (создан {recent_minutes}+ минут назад): {latest_file.name}")

        self.tool_files = collected_files
        return collected_files

    def merge_txt_reports_line_by_line(self, output_path: Optional[Path] = None, separator: str = '\n') -> str:
        """
        Объединяет TXT файлы из собранных отчетов в один результирующий файл

        Args:
            output_path: Путь к результирующему файлу (если None, создается автоматически)
            separator: Разделитель между файлами (по умолчанию '\n')

        Returns:
            str: Путь к созданному объединенному TXT файлу
        """
        if not self.tool_files:
            print("[!] Нет собранных файлов для объединения")
            return ""

        if output_path is None:
            output_path = self.txt_dir / f"{self.filename_base}.txt"
        else:
            output_path = Path(output_path)

        # Создаем директорию если нужно
        output_path.parent.mkdir(parents=True, exist_ok=True)
        # Установка прав доступа для директорий
        os.chmod(output_path.parent, 0o755)

        try:
            # Порядок инструментов для объединения
            tool_order = ['wappalyzer', 'scanner', 'nuclei',
                          'whois', 'web', 'osint', 'amass', 'katana', 'retire', 'cors', 'ssl-tls', 'dns', 'nmap']

            # Заголовок объединенного отчета
            with open(output_path, 'w', encoding='utf-8') as output_file:
                # Пишем заголовок
                header = f"""
╔{'═' * 78}╗
║{'ОБЪЕДИНЕННЫЙ ОТЧЕТ О СКАНИРОВАНИИ'.center(78)}║
╚{'═' * 78}╝

📋 ИНФОРМАЦИЯ О СКАНИРОВАНИИ
{'─' * 80}
  ID сканирования:      {self.scan_id}
  Дата создания:        {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}
  Использовано инструментов: {len(self.tool_files)}

{'═' * 80}

"""
                output_file.write(header)

                # Объединяем файлы
                total_lines = 0
                merged_tools = []

                for tool_name in tool_order:
                    if tool_name in self.tool_files:
                        file_path = self.tool_files[tool_name]
                        merged_tools.append(tool_name)
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                                lines = content.split('\n')
                                line_count = len(
                                    [l for l in lines if l.strip()])
                                total_lines += line_count

                                # Добавляем разделитель между файлами
                                output_file.write(f"\n{'=' * 80}\n")
                                output_file.write(
                                    f"🔧 ОТЧЕТ ИНСТРУМЕНТА: {tool_name.upper()}\n")
                                output_file.write(f"{'─' * 80}\n")
                                output_file.write(content)
                                output_file.write(f"\n{'─' * 80}\n")
                                output_file.write(
                                    f"📊 Статистика: {line_count} строк(и)\n")

                                if separator:
                                    output_file.write(separator)

                        except Exception as e:
                            print(f"[!] Ошибка при чтении {file_path}: {e}")
                            output_file.write(f"\n{'─' * 80}\n")
                            output_file.write(
                                f"⚠ ОШИБКА: Не удалось прочитать отчет {tool_name}\n")
                            output_file.write(f"   {str(e)}\n")

                # Пишем итоговую статистику
                footer = f"""
{'═' * 80}
📊 ИТОГОВАЯ СТАТИСТИКА
{'─' * 80}
  Всего обработано отчетов: {len(self.tool_files)}
  Объединенные инструменты: {', '.join(merged_tools)}
  Общее количество строк:    {total_lines}
{'═' * 80}

Дата создания отчета: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}
{'═' * 80}
"""
                output_file.write(footer)

            print(f"\n[+] TXT отчеты успешно объединены в {output_path}")
            print(f"[+] Статистика:")
            for tool_name in tool_order:
                if tool_name in self.tool_files:
                    file_path = self.tool_files[tool_name]
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            lines = len(
                                [l for l in f.readlines() if l.strip()])
                            print(f"    {tool_name:15} - {lines:5} строк")
                    except:
                        print(f"    {tool_name:15} - {'ошибка':>5}")
            print(f"    {'─' * 20}")
            print(f"    {'ВСЕГО':15} - {total_lines:5} строк")
            
            # Установка прав доступа для чтения всем пользователям
            os.chmod(output_path, 0o644)

            return str(output_path)

        except Exception as e:
            print(f"[!] Ошибка при объединении файлов: {e}")
            return ""

    def merge_txt_reports_with_structure(self, output_path: Optional[Path] = None) -> str:
        """
        Объединяет TXT файлы с сохранением структуры и добавлением разделителей

        Args:
            output_path: Путь к результирующему файлу

        Returns:
            str: Путь к созданному объединенному TXT файлу
        """
        if not self.tool_files:
            print("[!] Нет собранных файлов для объединения")
            return ""

        if output_path is None:
            output_path = self.txt_dir / f"{self.filename_base}.txt"
        else:
            output_path = Path(output_path)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        # Установка прав доступа для директорий
        os.chmod(output_path.parent, 0o755)

        try:
            with open(output_path, 'w', encoding='utf-8') as output_file:
                # Заголовок
                output_file.write(f"""
╔{'═' * 78}╗
║{'ОБЪЕДИНЕННЫЙ ОТЧЕТ О СКАНИРОВАНИИ'.center(78)}║
╚{'═' * 78}╝

📋 ИНФОРМАЦИЯ О СКАНИРОВАНИИ
{'─' * 80}
  ID сканирования:      {self.scan_id}
  Дата создания:        {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}
  Целевые адреса:       {', '.join(self.targets) if self.targets else 'Не определены'}
  Инструментов:         {len(self.tool_files)}

""")

                # Сортируем инструменты для последовательного вывода
                sorted_tools = sorted(self.tool_files.keys())

                for idx, (tool_name, file_path) in enumerate(sorted_tools, 1):
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()

                            # Добавляем разделитель и заголовок
                            output_file.write(f"\n{'═' * 80}\n")
                            output_file.write(
                                f"🔧 ИНСТРУМЕНТ {idx}: {tool_name.upper()}\n")
                            output_file.write(f"{'─' * 80}\n")
                            output_file.write(content)
                            output_file.write(f"\n{'─' * 80}\n")

                    except Exception as e:
                        output_file.write(f"\n{'─' * 80}\n")
                        output_file.write(
                            f"⚠ ОШИБКА: Не удалось прочитать отчет {tool_name}\n")
                        output_file.write(f"   {str(e)}\n")

                # Футер
                output_file.write(f"""
{'═' * 80}
📊 ИТОГОВАЯ СТАТИСТИКА
{'─' * 80}
  Всего обработано отчетов: {len(self.tool_files)}
{'═' * 80}

Дата создания отчета: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}
{'═' * 80}
""")

            print(f"[+] TXT отчеты успешно объединены в {output_path}")
            
            # Установка прав доступа для чтения всем пользователям
            os.chmod(output_path, 0o644)
            
            return str(output_path)

        except Exception as e:
            print(f"[!] Ошибка при объединении файлов: {e}")
            return ""

    def get_combined_json(self) -> Dict[str, Any]:
        """Получить объединенный отчет в JSON с полной информацией"""
        combined_json = {
            'scan_id': self.scan_id,
            'scan_datetime': self.scan_datetime,
            'targets': list(self.targets),
            'tools_count': len(self.tool_reports),
            'files_merged': {
                tool: str(file_path) for tool, file_path in self.tool_files.items()
            },
            'summary': self._generate_summary(),
            'detailed_findings': self._extract_detailed_findings(),
            'tool_reports': self.tool_reports
        }
        return combined_json

    def _extract_detailed_findings(self) -> Dict[str, Any]:
        """Извлечь подробные находки из всех инструментов"""
        detailed = {
            'vulnerabilities': [],
            'subdomains': [],
            'urls': [],
            'technologies': [],
            'hosts': [],
            'dns_records': [],
            'nameservers': []
        }

        for tool_name, report in self.tool_reports.items():
            # Уязвимости от Nuclei
            if tool_name == 'nuclei' and 'results' in report:
                for vuln in report.get('results', []):
                    detailed['vulnerabilities'].append({
                        'tool': 'nuclei',
                        'template_id': vuln.get('template_id'),
                        'name': vuln.get('info', {}).get('name'),
                        'severity': vuln.get('info', {}).get('severity'),
                        'url': vuln.get('matched_at'),
                        'description': vuln.get('info', {}).get('description')
                    })

            # Уязвимости от Scanner
            elif tool_name == 'scanner' and 'vulnerabilities' in report:
                for vuln in report.get('vulnerabilities', []):
                    detailed['vulnerabilities'].append({
                        'tool': 'scanner',
                        'type': vuln.get('type'),
                        'severity': vuln.get('severity'),
                        'url': vuln.get('url'),
                        'description': vuln.get('description')
                    })

            # Поддомены от Amass, Osint
            elif tool_name in ['amass', 'osint'] and 'subdomains' in report:
                for subdomain in report.get('subdomains', []):
                    if subdomain not in detailed['subdomains']:
                        detailed['subdomains'].append(subdomain)

            # Живые хосты
            if tool_name == 'osint' and 'alive_hosts' in report:
                for host in report.get('alive_hosts', []):
                    if host not in detailed['hosts']:
                        detailed['hosts'].append(host)

            # URL от Web сканнера
            elif tool_name == 'web' and 'urls' in report:
                for url in report.get('urls', []):
                    if url not in detailed['urls']:
                        detailed['urls'].append(url)

            # Уязвимости от Retire (JS библиотеки)
            elif tool_name == 'retire' and 'vulnerabilities' in report:
                for vuln in report.get('vulnerabilities', []):
                    detailed['vulnerabilities'].append({
                        'tool': 'retire',
                        'type': 'javascript_vulnerability',
                        'severity': vuln.get('severity', 'unknown'),
                        'id': vuln.get('id'),
                        'info': vuln.get('info'),
                        'description': vuln.get('info')
                    })

            # Технологии
            elif tool_name == 'wappalyzer' and 'technologies' in report:
                for tech in report.get('technologies', []):
                    tech_name = tech.get('name') if isinstance(
                        tech, dict) else str(tech)
                    if tech_name not in detailed['technologies']:
                        detailed['technologies'].append({
                            'name': tech_name,
                            'version': tech.get('version') if isinstance(tech, dict) else 'unknown'
                        })

            # DNS записи
            elif tool_name == 'dns' and 'dns_records' in report:
                # Добавляем nameservers
                for ns in report.get('nameservers', []):
                    if ns not in detailed['nameservers']:
                        detailed['nameservers'].append(ns)
                
                # Добавляем DNS записи
                for record_type, records in report.get('dns_records', {}).items():
                    for record in records:
                        if isinstance(record, dict):
                            # Для NS записей с IP-адресами
                            if 'nameserver' in record:
                                detailed['dns_records'].append({
                                    'type': 'NS',
                                    'nameserver': record.get('nameserver'),
                                    'ips': record.get('ips', [])
                                })
                            # Для MX записей
                            elif 'server' in record and 'priority' in record:
                                detailed['dns_records'].append({
                                    'type': 'MX',
                                    'server': record.get('server'),
                                    'priority': record.get('priority')
                                })
                            # Для SOA записей
                            elif 'mname' in record:
                                detailed['dns_records'].append({
                                    'type': 'SOA',
                                    'mname': record.get('mname'),
                                    'rname': record.get('rname'),
                                    'serial': record.get('serial')
                                })
                        else:
                            # Простые записи (A, AAAA, CNAME, TXT)
                            detailed['dns_records'].append({
                                'type': record_type,
                                'value': str(record)
                            })

            # Данные от Nmap
            elif tool_name == 'nmap':
                # Уязвимости от Nmap
                if 'vulnerabilities' in report:
                    for vuln in report.get('vulnerabilities', []):
                        detailed['vulnerabilities'].append({
                            'tool': 'nmap',
                            'cve_id': vuln.get('cve_id', 'Unknown'),
                            'severity': vuln.get('severity', 'unknown'),
                            'service': vuln.get('service'),
                            'port': vuln.get('port'),
                            'host': vuln.get('host'),
                            'version': vuln.get('version'),
                            'description': vuln.get('description')
                        })
                
                # Хосты и открытые порты от Nmap
                if 'hosts' in report:
                    for host in report.get('hosts', []):
                        host_addr = host.get('host')
                        if host_addr and host_addr not in detailed['hosts']:
                            detailed['hosts'].append(host_addr)
                
                # На структурированные данные
                if 'summary' in report and 'total_open_ports' in report['summary']:
                    # Добавляем подробные данные портов если доступны
                    if isinstance(report.get('hosts'), list):
                        for host_info in report.get('hosts', []):
                            if 'ports' in host_info:
                                for port in host_info.get('ports', []):
                                    if port.get('state') == 'open':
                                        detailed['urls'].append(
                                            f"{host_info.get('host')}:{port.get('port')}/{port.get('service')}"
                                        )

        return detailed

    def _generate_summary(self) -> Dict[str, Any]:
        """Генерировать краткое описание результатов"""
        summary = {
            'total_vulnerabilities': 0,
            'total_findings': 0,
            'by_severity': {
                'critical': 0,
                'high': 0,
                'medium': 0,
                'low': 0,
                'info': 0
            },
            'tools_executed': list(self.tool_reports.keys()),
            'tool_summaries': {}
        }

        # Подсчитываем данные по каждому инструменту
        for tool_name, report in self.tool_reports.items():
            tool_summary = {
                'status': report.get('status', 'unknown'),
                'findings': 0
            }

            # Обработка результатов в зависимости от инструмента
            if tool_name == 'nuclei':
                if 'summary' in report:
                    nuclei_summary = report['summary']
                    tool_summary['findings'] = nuclei_summary.get('total', 0)
                    summary['total_vulnerabilities'] += nuclei_summary.get(
                        'total', 0)

                    for severity in ['critical', 'high', 'medium', 'low', 'info']:
                        summary['by_severity'][severity] += nuclei_summary.get(
                            'by_severity', {}).get(severity, 0)
                elif 'results' in report:
                    tool_summary['findings'] = len(report.get('results', []))
                    summary['total_vulnerabilities'] += len(
                        report.get('results', []))

            elif tool_name == 'scanner':
                if isinstance(report.get('vulnerabilities'), list):
                    tool_summary['findings'] = len(
                        report.get('vulnerabilities', []))
                    summary['total_vulnerabilities'] += len(
                        report.get('vulnerabilities', []))

            elif tool_name == 'wappalyzer':
                if isinstance(report.get('technologies'), list):
                    tool_summary['findings'] = len(
                        report.get('technologies', []))

            elif tool_name == 'amass':
                if 'subdomains_count' in report:
                    tool_summary['findings'] = report.get(
                        'subdomains_count', 0)
                    summary['total_findings'] += report.get(
                        'subdomains_count', 0)

            elif tool_name == 'osint':
                if 'subdomains_count' in report:
                    tool_summary['findings'] = report.get(
                        'subdomains_count', 0)
                    summary['total_findings'] += report.get(
                        'subdomains_count', 0)

            elif tool_name == 'web':
                if 'urls' in report:
                    tool_summary['findings'] = len(report.get('urls', []))
                    summary['total_findings'] += len(report.get('urls', []))

            elif tool_name == 'whois':
                if 'output' in report and report['output']:
                    tool_summary['findings'] = 1

            elif tool_name == 'retire':
                vulns = report.get('vulnerabilities', [])
                tool_summary['findings'] = len(vulns)
                summary['total_vulnerabilities'] += len(vulns)
                summary['total_findings'] += len(vulns)
                
                # Подсчитываем по серьезности
                for vuln in vulns:
                    if isinstance(vuln, dict):
                        severity = vuln.get('severity', 'low').lower()
                        if severity in summary['by_severity']:
                            summary['by_severity'][severity] += 1
                        else:
                            summary['by_severity']['low'] += 1

            elif tool_name == 'dns':
                dns_records = report.get('dns_records', {})
                total_records = sum(len(records) for records in dns_records.values()) if isinstance(dns_records, dict) else 0
                nameservers = report.get('nameservers', [])
                tool_summary['findings'] = total_records
                tool_summary['dns_records_count'] = total_records
                tool_summary['nameservers_count'] = len(nameservers)
                summary['total_findings'] += total_records

            elif tool_name == 'nmap':
                # Обработка результатов Nmap
                if 'summary' in report:
                    summary_data = report['summary']
                    total_vulns = summary_data.get('total_vulnerabilities', 0)
                    total_ports = summary_data.get('total_open_ports', 0)
                    
                    tool_summary['findings'] = total_ports
                    tool_summary['vulnerabilities'] = total_vulns
                    tool_summary['hosts_discovered'] = summary_data.get('total_hosts_discovered', 0)
                    tool_summary['open_ports'] = total_ports
                    
                    summary['total_vulnerabilities'] += total_vulns
                    summary['total_findings'] += total_ports
                    
                    # Подсчитываем уязвимости по серьезности
                    for severity, count in summary_data.get('vulnerabilities_by_severity', {}).items():
                        severity_lower = severity.lower()
                        if severity_lower in summary['by_severity']:
                            summary['by_severity'][severity_lower] += count
                
                # Или процесс выше не сработал, используем уязвимости напрямую
                elif 'vulnerabilities' in report:
                    vulns = report.get('vulnerabilities', [])
                    tool_summary['findings'] = len(vulns) if isinstance(vulns, list) else 0
                    summary['total_vulnerabilities'] += tool_summary['findings']
                    
                    # Подсчитываем по серьезности
                    for vuln in vulns:
                        if isinstance(vuln, dict):
                            severity = vuln.get('severity', 'unknown').lower()
                            if severity in summary['by_severity']:
                                summary['by_severity'][severity] += 1

            summary['tool_summaries'][tool_name] = tool_summary

        return summary

    def save_json(self) -> str:
        """Сохранить объединенный JSON отчет"""
        data = self.get_combined_json()
        report_path = self.json_dir / f"{self.filename_base}.json"

        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        # Установка прав доступа для чтения всем пользователям
        os.chmod(report_path, 0o644)

        print(f"[+] Объединенный JSON отчет сохранен: {report_path}")
        return str(report_path)

    def save_txt(self, method: str = 'line_by_line') -> str:
        """
        Сохранить объединенный TXT отчет

        Args:
            method: Метод объединения ('line_by_line' или 'structured')
        """
        if method == 'line_by_line':
            return self.merge_txt_reports_line_by_line()
        else:
            return self.merge_txt_reports_with_structure()

    def txt_to_docx(self, txt_file_path: Optional[str] = None, output_path: Optional[Path] = None) -> str:
        """
        Преобразовать TXT отчет в DOCX (Word)
        Оптимизирована для больших файлов с автоматическим выбором стратегии

        Args:
            txt_file_path: Путь к TXT файлу (если None, используется последний созданный)
            output_path: Путь к выходному DOCX файлу

        Returns:
            str: Путь к созданному DOCX файлу
        """
        if not DOCX_AVAILABLE:
            print("[!] Модуль python-docx не установлен. Установите: pip install python-docx")
            return ""

        # Определяем путь к TXT файлу
        if txt_file_path is None:
            txt_file_path = self.txt_dir / f"{self.filename_base}.txt"
        else:
            txt_file_path = Path(txt_file_path)

        if not txt_file_path.exists():
            print(f"[!] TXT файл не найден: {txt_file_path}")
            return ""

        # Проверяем размер файла
        file_size_kb = txt_file_path.stat().st_size / 1024
        file_size_mb = file_size_kb / 1024
        
        if file_size_mb > 10:
            print(f"[*] Большой файл ({file_size_mb:.2f} MB). Обработка может занять время...")
        elif file_size_mb > 50:
            print(f"[*] ОЧЕНЬ большой файл ({file_size_mb:.2f} MB). Используем оптимизированную обработку...")

        # Определяем путь к выходному файлу
        if output_path is None:
            output_path = self.word_dir / f"{self.filename_base}.docx"
        else:
            output_path = Path(output_path)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        # Установка прав доступа для директорий
        os.chmod(output_path.parent, 0o755)

        try:
            print(f"[*] Начинаю преобразование TXT в DOCX: {txt_file_path.name} ({file_size_mb:.2f} MB)")
            doc = Document()

            # Читаем TXT файл и преобразуем в DOCX
            with open(txt_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            print(f"[*] TXT файл загружен ({len(content)} символов). Начинаю парсинг...")

            # Выбираем стратегию обработки на основе размера
            if file_size_mb > 100:
                print(f"[*] Использую режим таблиц для очень больших файлов")
                _parse_and_add_to_docx_table(doc, content)
            else:
                # Парсим содержимое и добавляем в документ
                _parse_and_add_to_docx(doc, content)

            # Сохраняем документ
            print(f"[*] Сохраняю DOCX документ...")
            doc.save(str(output_path))
            
            # Установка прав доступа для чтения всем пользователям
            os.chmod(output_path, 0o644)
            
            output_size_mb = output_path.stat().st_size / (1024 * 1024)
            print(f"[+] DOCX отчет успешно сохранен!")
            print(f"[+] Путь: {output_path}")
            print(f"[+] Размер DOCX: {output_size_mb:.2f} MB")
            return str(output_path)

        except MemoryError as e:
            print(f"[!] ОШИБКА ПАМЯТИ при преобразовании в DOCX (файл слишком большой)")
            print(f"[*] Файл: {txt_file_path.name} ({file_size_mb:.2f} MB)")
            print(f"[*] Рекомендации:")
            print(f"   - Закройте другие программы для освобождения памяти")
            print(f"   - Установите дополнительную виртуальную память (swap)")
            print(f"   - Разбейте файл на несколько частей вручную")
            return ""
        except Exception as e:
            print(f"[!] ОШИБКА при преобразовании в DOCX")
            print(f"[!] Тип: {type(e).__name__}")
            print(f"[!] Текст: {e}")
            print(f"[*] Файл: {txt_file_path.name}")
            import traceback
            print("[*] Полный стек ошибки:")
            traceback.print_exc()
            return ""

    def collect_and_merge_recent_with_docx(self, recent_minutes: int = 5, method: str = 'line_by_line', 
                                           include_docx: bool = True) -> Dict[str, str]:
        """
        Собрать недавние файлы, объединить их и создать DOCX отчет

        Args:
            recent_minutes: Количество минут для поиска недавних отчетов
            method: Метод объединения TXT ('line_by_line' или 'structured')
            include_docx: Создать ли DOCX версию отчета

        Returns:
            Dict с путями к JSON, TXT и DOCX отчетам
        """
        # Собираем недавние файлы
        self.collect_recent_files(recent_minutes=recent_minutes)

        # Сохраняем JSON
        json_path = self.save_json()

        # Сохраняем TXT
        txt_path = self.save_txt(method=method)

        # Сохраняем DOCX
        docx_path = ""
        if include_docx:
            docx_path = self.txt_to_docx(txt_path)

        return {
            'json': json_path,
            'txt': txt_path,
            'docx': docx_path
        }

    def collect_and_merge_recent(self, recent_minutes: int = 5, method: str = 'line_by_line') -> Dict[str, str]:
        """
        Собрать недавние файлы и объединить их в один отчет

        Args:
            recent_minutes: Количество минут для поиска недавних отчетов
            method: Метод объединения TXT ('line_by_line' или 'structured')

        Returns:
            Dict с путями к JSON и TXT отчетам
        """
        # Собираем недавние файлы
        self.collect_recent_files(recent_minutes=recent_minutes)

        # Сохраняем JSON
        json_path = self.save_json()

        # Сохраняем TXT
        txt_path = self.save_txt(method=method)

        return {
            'json': json_path,
            'txt': txt_path
        }

    def collect_and_merge_by_time_window(self, start_time: datetime,
                                         end_time: Optional[datetime] = None,
                                         method: str = 'line_by_line') -> Dict[str, str]:
        """
        Собрать файлы в указанном временном окне и объединить их

        Args:
            start_time: Начало временного окна
            end_time: Конец временного окна (если None, используется текущее время)
            method: Метод объединения TXT

        Returns:
            Dict с путями к JSON и TXT отчетам
        """
        # Собираем файлы во временном окне
        self.collect_files_by_time_window(
            start_time=start_time, end_time=end_time)

        # Сохраняем JSON
        json_path = self.save_json()

        # Сохраняем TXT
        txt_path = self.save_txt(method=method)

        return {
            'json': json_path,
            'txt': txt_path
        }


def create_combined_report(scan_id: str, reports_base_dir: Optional[Path] = None,
                           recent_minutes: int = 5, method: str = 'line_by_line',
                           include_docx: bool = True) -> Dict[str, str]:
    """
    Создать объединённый отчет из недавно созданных отчетов инструментов

    Args:
        scan_id: ID сканирования (уникальный идентификатор)
        reports_base_dir: Базовая директория для отчетов
        recent_minutes: Количество минут для поиска недавних отчетов (по умолчанию 5)
        method: Метод объединения TXT ('line_by_line' или 'structured')
        include_docx: Создать ли DOCX версию отчета

    Returns:
        Dict с путями к JSON, TXT и DOCX отчетам
    """
    combined = CombinedReport(scan_id, reports_base_dir)
    if include_docx:
        return combined.collect_and_merge_recent_with_docx(recent_minutes=recent_minutes, method=method)
    else:
        result = combined.collect_and_merge_recent(recent_minutes=recent_minutes, method=method)
        result['docx'] = ""
        return result


def create_combined_report_by_time(scan_id: str, start_time: datetime,
                                   end_time: Optional[datetime] = None,
                                   reports_base_dir: Optional[Path] = None,
                                   method: str = 'line_by_line',
                                   include_docx: bool = True) -> Dict[str, str]:
    """
    Создать объединённый отчет из отчетов, созданных в указанном временном окне

    Args:
        scan_id: ID сканирования
        start_time: Начало временного окна
        end_time: Конец временного окна (если None, используется текущее время)
        reports_base_dir: Базовая директория для отчетов
        method: Метод объединения TXT
        include_docx: Создать ли DOCX версию отчета

    Returns:
        Dict с путями к JSON, TXT и DOCX отчетам
    """
    combined = CombinedReport(scan_id, reports_base_dir)
    combined.collect_files_by_time_window(start_time=start_time, end_time=end_time)

    json_path = combined.save_json()
    txt_path = combined.save_txt(method=method)
    docx_path = ""
    
    if include_docx:
        docx_path = combined.txt_to_docx(txt_path)

    return {
        'json': json_path,
        'txt': txt_path,
        'docx': docx_path
    }


def quick_merge_all_reports(reports_base_dir: Optional[Path] = None,
                            output_name: str = "merged_report",
                            include_docx: bool = True) -> Dict[str, str]:
    """
    Быстрое объединение всех отчетов (без ограничений по времени)

    Args:
        reports_base_dir: Базовая директория отчетов
        output_name: Имя выходного файла (без расширения)
        include_docx: Создать ли DOCX версию отчета

    Returns:
        Dict с путями к JSON, TXT и DOCX отчетам
    """
    scan_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    combined = CombinedReport(scan_id, reports_base_dir)

    # Собираем все TXT файлы (без ограничения по времени)
    combined.collect_recent_files(recent_minutes=999999)  # Все файлы

    # Сохраняем с кастомным именем
    combined.filename_base = output_name

    json_path = combined.save_json()
    txt_path = combined.save_txt(method='line_by_line')
    docx_path = ""
    
    if include_docx:
        docx_path = combined.txt_to_docx(txt_path)

    return {
        'json': json_path,
        'txt': txt_path,
        'docx': docx_path
    }


def _parse_and_add_to_docx(doc: 'Document', content: str) -> None:
    """
    Парсить TXT содержимое и добавить его в DOCX документ с форматированием
    Оптимизирована для больших файлов

    Args:
        doc: Document объект из python-docx
        content: Текстовое содержимое для добавления
    """
    lines = content.split('\n')
    paragraph_count = 0
    max_paragraphs_per_batch = 100  # Обработка текста батчами

    try:
        for i, line in enumerate(lines):
            try:
                if not line.strip():
                    # Пустая строка
                    doc.add_paragraph()
                    paragraph_count += 1
                    continue

                # Определяем тип строки и добавляем с соответствующим форматированием
                if '╔' in line or '═' in line or '║' in line or '╚' in line:
                    # Это линии разделители - пропускаем (будут пустые строки)
                    continue
                elif line.startswith('ОБЪЕДИНЕННЫЙ ОТЧЕТ') or line.startswith('ИНФОРМАЦИЯ О СКАНИРОВАНИИ') or \
                     line.startswith('ИТОГОВАЯ СТАТИСТИКА') or line.startswith('ИНСТРУМЕНТ'):
                    # Заголовки разделов
                    p = doc.add_paragraph(line.strip())
                    p.style = 'Heading 1'
                    p_format = p.paragraph_format
                    p_format.space_before = Pt(12)
                    p_format.space_after = Pt(6)
                    paragraph_count += 1
                elif line.startswith('🔧') or line.startswith('📋') or line.startswith('📊'):
                    # Подзаголовки с иконками
                    p = doc.add_paragraph(line.strip())
                    p.style = 'Heading 2'
                    p_format = p.paragraph_format
                    p_format.space_before = Pt(6)
                    p_format.space_after = Pt(4)
                    paragraph_count += 1
                elif line.startswith('  '):
                    # Отступанные строки (параметры, данные)
                    p = doc.add_paragraph(line.lstrip(), style='List Bullet')
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.5)
                    paragraph_count += 1
                elif '─' in line:
                    # Разделители - пустая строка
                    continue
                else:
                    # Обычный текст
                    p = doc.add_paragraph(line)
                    p_format = p.paragraph_format
                    p_format.space_after = Pt(6)
                    paragraph_count += 1

            except Exception as line_error:
                # Логируем ошибку для конкретной строки, но продолжаем обработку
                print(f"[!] Ошибка при обработке строки {i+1}: {line_error}")
                # Добавляем строку как простой текст без форматирования
                try:
                    doc.add_paragraph(line[:500])  # Ограничиваем длину строки
                    paragraph_count += 1
                except:
                    pass  # Если не получилось даже обычный текст, пропускаем

            # Периодический логинг прогресса для больших файлов
            if paragraph_count % max_paragraphs_per_batch == 0 and paragraph_count > 0:
                print(f"[*] Обработано {paragraph_count} параграфов...")

    except Exception as e:
        print(f"[!] Критическая ошибка при парсинге DOCX: {e}")
        print(f"[*] Обработано {paragraph_count} параграфов перед ошибкой")
        raise


def _parse_and_add_to_docx_table(doc: 'Document', content: str) -> None:
    """
    Альтернативный парсинг с использованием таблиц - для очень больших файлов
    Более эффективен для файлов > 100 MB

    Args:
        doc: Document объект из python-docx
        content: Текстовое содержимое для добавления
    """
    lines = content.split('\n')
    section_lines = []
    current_section = None
    table = None

    try:
        for i, line in enumerate(lines):
            try:
                # Определяем заголовки разделов
                if line.startswith('🔧') or line.startswith('📋') or line.startswith('════'):
                    # Сохраняем предыдущую таблицу если есть
                    if section_lines and table is None:
                        # Создаем таблицу для секции
                        table = doc.add_table(rows=len(section_lines) + 1, cols=1)
                        table.style = 'Light Grid Accent 1'
                        cell = table.rows[0].cells[0]
                        cell.text = current_section or "Содержимое"
                        for idx, sec_line in enumerate(section_lines):
                            cell = table.rows[idx + 1].cells[0]
                            cell.text = sec_line[:200]  # Ограничиваем длину

                    # Начинаем новую секцию
                    doc.add_paragraph()
                    p = doc.add_paragraph(line.strip())
                    if '🔧' in line or '📋' in line:
                        p.style = 'Heading 2'
                    current_section = line.strip()
                    section_lines = []
                    table = None

                elif not line.strip():
                    continue

                elif '═' in line or '║' in line or '╔' in line or '╚' in line or '─' in line:
                    continue

                else:
                    section_lines.append(line[:200])  # Ограничиваем строки

                    # Если секция слишком большая, сохраняем в таблицу
                    if len(section_lines) >= 50:
                        table = doc.add_table(rows=len(section_lines) + 1, cols=1)
                        table.style = 'Light Grid Accent 1'
                        cell = table.rows[0].cells[0]
                        cell.text = current_section or "Содержимое"
                        for idx, sec_line in enumerate(section_lines):
                            cell = table.rows[idx + 1].cells[0]
                            cell.text = sec_line
                        section_lines = []

            except Exception as line_error:
                print(f"[!] Ошибка на строке {i+1}: {line_error}")
                continue

    except Exception as e:
        print(f"[!] Ошибка при парсинге с таблицами: {e}")
        raise


def txt_to_docx_file(txt_path: str, output_path: Optional[str] = None) -> str:
    """
    Преобразовать отдельный TXT файл в DOCX
    Оптимизирована для больших файлов с автоматическим выбором стратегии

    Args:
        txt_path: Путь к TXT файлу
        output_path: Путь к выходному DOCX файлу (если None, создается рядом с TXT)

    Returns:
        str: Путь к созданному DOCX файлу
    """
    if not DOCX_AVAILABLE:
        print("[!] Модуль python-docx не установлен. Установите: pip install python-docx")
        return ""

    txt_path = Path(txt_path)
    if not txt_path.exists():
        print(f"[!] TXT файл не найден: {txt_path}")
        return ""

    # Проверяем размер файла
    file_size_kb = txt_path.stat().st_size / 1024
    file_size_mb = file_size_kb / 1024
    
    if file_size_mb > 10:
        print(f"[*] Большой файл ({file_size_mb:.2f} MB). Обработка может занять время...")
    elif file_size_mb > 50:
        print(f"[*] ОЧЕНЬ большой файл ({file_size_mb:.2f} MB). Используем оптимизированную обработку...")

    # Определяем выходной путь
    if output_path is None:
        output_path = txt_path.parent / f"{txt_path.stem}.docx"
    else:
        output_path = Path(output_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    # Установка прав доступа для директорий
    os.chmod(output_path.parent, 0o755)

    try:
        print(f"[*] Начинаю преобразование TXT в DOCX: {txt_path.name} ({file_size_mb:.2f} MB)")
        doc = Document()

        # Читаем TXT файл
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        print(f"[*] TXT файл загружен ({len(content)} символов). Начинаю парсинг...")

        # Выбираем стратегию обработки на основе размера
        if file_size_mb > 100:
            print(f"[*] Использую режим таблиц для очень больших файлов")
            _parse_and_add_to_docx_table(doc, content)
        else:
            # Парсим и добавляем в документ
            _parse_and_add_to_docx(doc, content)

        # Сохраняем
        print(f"[*] Сохраняю DOCX документ...")
        doc.save(str(output_path))
        
        # Установка прав доступа для чтения всем пользователям
        os.chmod(output_path, 0o644)
        
        output_size_mb = output_path.stat().st_size / (1024 * 1024)
        print(f"[+] DOCX файл успешно сохранен!")
        print(f"[+] Путь: {output_path}")
        print(f"[+] Размер DOCX: {output_size_mb:.2f} MB")
        return str(output_path)

    except MemoryError as e:
        print(f"[!] ОШИБКА ПАМЯТИ при преобразовании в DOCX (файл слишком большой)")
        print(f"[*] Файл: {txt_path.name} ({file_size_mb:.2f} MB)")
        print(f"[*] Рекомендации:")
        print(f"   - Закройте другие программы для освобождения памяти")
        print(f"   - Установите дополнительную виртуальную память (swap)")
        print(f"   - Разбейте файл на несколько частей вручную")
        return ""
    except Exception as e:
        print(f"[!] ОШИБКА при преобразовании в DOCX")
        print(f"[!] Тип: {type(e).__name__}")
        print(f"[!] Текст: {e}")
        print(f"[*] Файл: {txt_path.name}")
        import traceback
        print("[*] Полный стек ошибки:")
        traceback.print_exc()
        return ""


# Пример использования

if __name__ == "__main__":

    # Пример 1️⃣: Создание объединенного отчета из недавних отчетов (последние 5 минут) с DOCX
    """
    scan_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    result = create_combined_report(scan_id, recent_minutes=5, method='line_by_line', include_docx=True)
    print(f"\n✅ Объединенный отчет создан:")
    print(f"   JSON: {result['json']}")
    print(f"   TXT:  {result['txt']}")
    print(f"   DOCX: {result['docx']}")
    """

    # Пример 2️⃣: Создание отчета с указанием временного окна с DOCX
    scan_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    start_time = datetime.now() - timedelta(minutes=10)
    result2 = create_combined_report_by_time(
        scan_id, start_time, method='line_by_line', include_docx=True)
    print(f"\n✅ Объединенный отчет создан:")
    print(f"   JSON: {result2['json']}")
    print(f"   TXT:  {result2['txt']}")
    if result2['docx']:
        print(f"   DOCX: {result2['docx']}")

    # Пример 3️⃣: Быстрое объединение всех отчетов с DOCX
    """
    result3 = quick_merge_all_reports(output_name="all_reports_merged", include_docx=True)
    print(f"\n✅ Все отчеты объединены:")
    print(f"   JSON: {result3['json']}")
    print(f"   TXT:  {result3['txt']}")
    print(f"   DOCX: {result3['docx']}")
    """

    # Пример 4️⃣: Преобразование отдельного TXT файла в DOCX
    """
    docx_path = txt_to_docx_file('/path/to/report.txt')
    print(f"✅ TXT преобразован в DOCX: {docx_path}")
    """
