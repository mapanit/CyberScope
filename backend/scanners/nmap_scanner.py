#!/usr/bin/env python3
"""
Nmap Port Scanner - Расширенное сетевое сканирование с поддержкой веб-приложений
Поддержка URL (http/https) и IP адресов
Отчеты в JSON, TXT и WORD форматах (БЕЗ КЭШИРОВАНИЯ В БД)
"""

import json
import requests
import sys
import os
import time
import argparse
import re
import ssl
import socket
from pathlib import Path
from datetime import datetime, timedelta
from urllib.parse import quote, urlparse
from typing import Dict, List, Optional, Any
from colorama import Fore, Style, init
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from dataclasses import dataclass
from enum import Enum

try:
    from core.report_utils import ReportBase
except ImportError:
    ReportBase = object

try:
    import nmap
    HAS_NMAP = True
except ImportError:
    HAS_NMAP = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Импорт конфигурации сканеров
try:
    from scanner_config import get_nmap_profile
except ImportError:
    get_nmap_profile = None

init(autoreset=True)

MAX_CVE_PER_SERVICE = 10
REQUEST_TIMEOUT = 10
MAX_WORKERS = 10
RATE_LIMIT_LOCK = threading.Lock()

class Severity(Enum):
    CRITICAL = "Critical"
    HIGH = "High"
    MEDIUM = "Medium"
    LOW = "Low"
    UNKNOWN = "Unknown"

@dataclass
class CVEInfo:
    """Информация о CVE уязвимости"""
    cve_id: str
    severity: str
    cvss_score: float
    description: str
    published_date: str
    exploit_maturity: str = "Unknown"
    remediation: str = ""
    references: List[str] = None
    
    def __post_init__(self):
        if self.references is None:
            self.references = []

class CVESearcher:
    """Расширенный поиск CVE - БЕЗ КЭШИРОВАНИЯ"""
    
    def __init__(self, nvd_api_key: str = None):
        self.nvd_api_key = nvd_api_key or os.environ.get('NVD_API_KEY')
        self.request_count = 0
        self.last_request_time = 0
        
    def _rate_limit(self):
        """Ограничение частоты запросов к API"""
        with RATE_LIMIT_LOCK:
            self.request_count += 1
            current_time = time.time()
            
            if not self.nvd_api_key:
                if self.request_count >= 5:
                    time_since_last = current_time - self.last_request_time
                    if time_since_last < 30:
                        time.sleep(30 - time_since_last)
                    self.request_count = 0
            else:
                if self.request_count >= 50:
                    time_since_last = current_time - self.last_request_time
                    if time_since_last < 30:
                        time.sleep(30 - time_since_last)
                    self.request_count = 0
            
            self.last_request_time = time.time()
    
    def search_cve(self, product: str, version: str = None) -> List[CVEInfo]:
        """Поиск CVE для продукта"""
        results = []
        
        nvd_results = self._search_nvd_api(product, version)
        results.extend(nvd_results)
        
        if len(results) < MAX_CVE_PER_SERVICE:
            circl_results = self._search_circl_api(product, version)
            results.extend(circl_results)
        
        local_results = self._search_known_vulnerabilities(product, version)
        results.extend(local_results)
        
        unique_results = {}
        for cve in results:
            if cve.cve_id not in unique_results:
                unique_results[cve.cve_id] = cve
        
        sorted_results = sorted(unique_results.values(), key=lambda x: x.cvss_score, reverse=True)
        return sorted_results[:MAX_CVE_PER_SERVICE]
    
    def _search_nvd_api(self, product: str, version: str = None) -> List[CVEInfo]:
        """Поиск через официальное NVD API"""
        try:
            base_url = "https://services.nvd.nist.gov/rest/json/cves/2.0"
            
            if version:
                keyword = f'"{product}" "{version}"'
            else:
                keyword = f'"{product}"'
            
            params = {
                'keywordSearch': keyword,
                'resultsPerPage': MAX_CVE_PER_SERVICE,
                'pubStartDate': (datetime.now() - timedelta(days=365)).isoformat()
            }
            
            if self.nvd_api_key:
                params['apiKey'] = self.nvd_api_key
            
            self._rate_limit()
            response = requests.get(base_url, params=params, timeout=REQUEST_TIMEOUT)
            
            if response.status_code == 200:
                data = response.json()
                cves = []
                
                for vuln in data.get('vulnerabilities', []):
                    cve = vuln.get('cve', {})
                    metrics = cve.get('metrics', {})
                    cvss_v3 = metrics.get('cvssMetricV31', [{}])[0].get('cvssData', {})
                    if not cvss_v3:
                        cvss_v3 = metrics.get('cvssMetricV30', [{}])[0].get('cvssData', {})
                    
                    severity = cvss_v3.get('baseSeverity', 'UNKNOWN')
                    cvss_score = cvss_v3.get('baseScore', 0.0)
                    
                    cves.append(CVEInfo(
                        cve_id=cve.get('id', 'Unknown'),
                        severity=severity.capitalize() if severity != 'UNKNOWN' else 'Unknown',
                        cvss_score=cvss_score,
                        description=cve.get('descriptions', [{}])[0].get('value', 'No description'),
                        published_date=cve.get('published', ''),
                        references=[ref.get('url', '') for ref in cve.get('references', [])[:3]]
                    ))
                
                return cves
        except Exception as e:
            pass
        
        return []
    
    def _search_circl_api(self, product: str, version: str = None) -> List[CVEInfo]:
        """Поиск через CIRCL CVE Search API"""
        try:
            base_url = "https://cve.circl.lu/api/search"
            query = f"{product} {version}" if version else product
            
            self._rate_limit()
            response = requests.get(f"{base_url}/{quote(query)}", timeout=REQUEST_TIMEOUT)
            
            if response.status_code == 200:
                data = response.json()
                cves = []
                
                for cve in data[:MAX_CVE_PER_SERVICE]:
                    cvss_score = cve.get('cvss', 0)
                    severity = "Critical" if cvss_score >= 9.0 else "High" if cvss_score >= 7.0 else "Medium" if cvss_score >= 4.0 else "Low"
                    
                    cves.append(CVEInfo(
                        cve_id=cve.get('id', 'Unknown'),
                        severity=severity,
                        cvss_score=cvss_score,
                        description=cve.get('summary', 'No description'),
                        published_date=cve.get('Published', ''),
                        references=cve.get('references', [])
                    ))
                
                return cves
        except Exception as e:
            pass
        
        return []
    
    def _search_known_vulnerabilities(self, product: str, version: str = None) -> List[CVEInfo]:
        """Локальная база известных уязвимостей"""
        known_vulns = {
            'openssl': [
                CVEInfo(
                    cve_id='CVE-2014-0160',
                    severity='Critical',
                    cvss_score=10.0,
                    description='Heartbleed - Buffer over-read in OpenSSL',
                    published_date='2014-04-07',
                    exploit_maturity='High',
                    remediation='Обновите OpenSSL до версии 1.0.1g или выше'
                )
            ],
            'apache': [
                CVEInfo(
                    cve_id='CVE-2021-41773',
                    severity='Critical',
                    cvss_score=9.8,
                    description='Path traversal and RCE in Apache HTTP Server 2.4.49',
                    published_date='2021-10-05',
                    exploit_maturity='High',
                    remediation='Обновите Apache до версии 2.4.50'
                )
            ],
            'nginx': [
                CVEInfo(
                    cve_id='CVE-2021-23017',
                    severity='High',
                    cvss_score=8.6,
                    description='DNS resolver off-by-one error',
                    published_date='2021-08-18',
                    exploit_maturity='Medium',
                    remediation='Обновите nginx до версии 1.21.1'
                )
            ]
        }
        
        product_lower = product.lower()
        results = []
        
        for vuln_product, cves in known_vulns.items():
            if vuln_product in product_lower:
                results.extend(cves)
        
        return results

class NmapScanner(ReportBase):
    """Расширенный сканер с поддержкой IP адресов и URL веб-приложений"""
    
    def __init__(self, target: str, output_base: str = None, reports_dir: str = None, nvd_api_key: str = None):
        """Инициализация сканера"""
        super().__init__('nmap', target, Path(reports_dir) if reports_dir else None)
        
        self.target = target
        self.is_url = target.startswith(('http://', 'https://'))
        self.is_ip = self._is_valid_ip_or_domain(target)
        
        self.discovered_hosts = []
        self.open_ports = []
        self.vulnerabilities = []
        self.web_vulnerabilities = []
        self.nse_results = []
        self.web_info = {}
        self.scan_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.start_time = datetime.now()
        
        self.cve_searcher = CVESearcher(nvd_api_key)
        self.reports_dir = self.tool_dir.parent
        
        if output_base:
            self.filename_base = output_base
        else:
            safe_target = target.replace('/', '_').replace('\\', '_').replace(':', '_').replace('?', '_')
            self.filename_base = f"nmap_{safe_target}_{self.scan_id}"
        
        self.scan_profile = "vuln"
        
        self.word_dir = self.reports_dir / "word"
        self.word_dir.mkdir(parents=True, exist_ok=True)
        
        # Для nmap сканирования
        if HAS_NMAP:
            self.nm = nmap.PortScanner()
        else:
            self.nm = None
        
        print(f"{Fore.GREEN}[*] Папки для отчетов:")
        print(f"{Fore.GREEN}[*]   JSON: {self.json_dir}")
        print(f"{Fore.GREEN}[*]   TXT: {self.txt_dir}")
        print(f"{Fore.GREEN}[*]   WORD: {self.word_dir}")
        print(f"{Fore.GREEN}[*] Тип цели: {'URL' if self.is_url else 'IP/Домен'}")
    
    def _is_valid_ip_or_domain(self, target: str) -> bool:
        """Проверить валидность IP или домена"""
        if target.startswith(('http://', 'https://')):
            return True
        
        # Простая проверка IP
        try:
            socket.inet_aton(target.split(':')[0])
            return True
        except socket.error:
            pass
        
        # Проверка домена
        domain_pattern = r'^[a-zA-Z0-9]([a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(\.[a-zA-Z0-9]([a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)*$'
        return bool(re.match(domain_pattern, target.split(':')[0]))
    
    def _scan_web_application(self) -> bool:
        """Сканирование веб-приложения через URL"""
        print(f"{Fore.CYAN}[*] Сканирование веб-приложения: {self.target}")
        
        try:
            # Парсим URL
            parsed_url = urlparse(self.target)
            host = parsed_url.hostname or self.target
            port = parsed_url.port or (443 if parsed_url.scheme == 'https' else 80)
            scheme = parsed_url.scheme or 'http'
            
            self.web_info = {
                'url': self.target,
                'host': host,
                'port': port,
                'scheme': scheme,
                'technologies': [],
                'headers': {},
                'status_code': None,
                'server': None,
                'powered_by': None,
                'issues': []
            }
            
            # Проверяем соединение и получаем информацию
            self._check_web_server()
            self._check_ssl_certificate()
            self._check_security_headers()
            self._check_common_vulnerabilities()
            
            print(f"{Fore.GREEN}[+] Веб-сканирование завершено")
            return True
            
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка веб-сканирования: {e}")
            self.web_vulnerabilities.append({
                'type': 'Connection Error',
                'severity': 'High',
                'description': str(e),
                'host': self.target
            })
            return False
    
    def _check_web_server(self):
        """Проверка веб-сервера"""
        try:
            response = requests.head(
                self.target,
                timeout=10,
                allow_redirects=True,
                verify=False
            )
            
            self.web_info['status_code'] = response.status_code
            headers = response.headers
            
            self.web_info['server'] = headers.get('Server', 'Unknown')
            self.web_info['powered_by'] = headers.get('X-Powered-By', None)
            self.web_info['headers'] = dict(headers)
            
            # Определяем технологии
            if 'Apache' in self.web_info['server']:
                self.web_info['technologies'].append('Apache')
            if 'nginx' in self.web_info['server']:
                self.web_info['technologies'].append('Nginx')
            if 'IIS' in self.web_info['server']:
                self.web_info['technologies'].append('IIS')
            if 'Powered-by' in headers and 'PHP' in headers.get('X-Powered-By', ''):
                self.web_info['technologies'].append('PHP')
            
        except Exception as e:
            print(f"{Fore.YELLOW}[!] Ошибка проверки сервера: {e}")
    
    def _check_ssl_certificate(self):
        """Проверка SSL сертификата"""
        if not self.target.startswith('https'):
            return
        
        try:
            parsed_url = urlparse(self.target)
            hostname = parsed_url.hostname
            port = parsed_url.port or 443
            
            context = ssl.create_default_context()
            context.check_hostname = False
            context.verify_mode = ssl.CERT_NONE
            
            with socket.create_connection((hostname, port), timeout=10) as sock:
                with context.wrap_socket(sock, server_hostname=hostname) as ssock:
                    cert = ssock.getpeercert()
                    
                    if cert:
                        # Проверка срока действия
                        not_after = cert.get('notAfter')
                        if not_after:
                            self.web_info['ssl_expires'] = not_after
                            
                            try:
                                from ssl import cert_time_to_seconds
                                expire_time = cert_time_to_seconds(not_after)
                                days_left = (expire_time - time.time()) / 86400
                                
                                if days_left < 0:
                                    self.web_vulnerabilities.append({
                                        'type': 'SSL Certificate Expired',
                                        'severity': 'Critical',
                                        'description': f'SSL сертификат истек {abs(int(days_left))} дней назад',
                                        'host': self.target
                                    })
                                elif days_left < 30:
                                    self.web_vulnerabilities.append({
                                        'type': 'SSL Certificate Expires Soon',
                                        'severity': 'High',
                                        'description': f'SSL сертификат истечет через {int(days_left)} дней',
                                        'host': self.target
                                    })
                            except:
                                pass
        except Exception as e:
            print(f"{Fore.YELLOW}[!] Ошибка проверки SSL: {e}")
    
    def _check_security_headers(self):
        """Проверка безопасности заголовков"""
        headers_to_check = {
            'X-Frame-Options': 'Защита от Clickjacking',
            'X-Content-Type-Options': 'Защита от MIME sniffing',
            'Content-Security-Policy': 'Политика безопасности контента',
            'Strict-Transport-Security': 'Принудительное HTTPS',
            'X-XSS-Protection': 'Защита от XSS'
        }
        
        try:
            response = requests.head(
                self.target,
                timeout=10,
                allow_redirects=True,
                verify=False
            )
            
            for header, description in headers_to_check.items():
                if header not in response.headers:
                    self.web_vulnerabilities.append({
                        'type': f'Missing {header}',
                        'severity': 'Medium',
                        'description': f'Отсутствует заголовок: {description}',
                        'header': header,
                        'host': self.target
                    })
        except Exception as e:
            print(f"{Fore.YELLOW}[!] Ошибка проверки заголовков: {e}")
    
    def _check_common_vulnerabilities(self):
        """Проверка общих уязвимостей"""
        common_urls = [
            '/admin/',
            '/.git/',
            '/.env',
            '/backup/',
            '/config/',
            '/web.config',
            '/xmlrpc.php',
            '/.htaccess',
            '/wp-admin/',
            '/.well-known/'
        ]
        
        try:
            for url_path in common_urls:
                try:
                    response = requests.get(
                        self.target.rstrip('/') + url_path,
                        timeout=5,
                        verify=False,
                        allow_redirects=False
                    )
                    
                    if response.status_code in [200, 301, 302]:
                        severity = 'Critical' if url_path in ['/.env', '/web.config'] else 'High'
                        self.web_vulnerabilities.append({
                            'type': 'Exposed Directory/File',
                            'severity': severity,
                            'description': f'Обнаружена открытая директория/файл: {url_path} (HTTP {response.status_code})',
                            'url': self.target.rstrip('/') + url_path,
                            'host': self.target
                        })
                except requests.Timeout:
                    pass
                except:
                    pass
        except Exception as e:
            print(f"{Fore.YELLOW}[!] Ошибка проверки уязвимостей: {e}")
    
    def run_scan(self, arguments: str = None, scan_profile: str = "vuln") -> bool:
        """Запуск сканирования (URL или IP)"""
        self.scan_profile = scan_profile
        
        # Если это URL, сканируем как веб-приложение
        if self.is_url:
            return self._scan_web_application()
        
        # Иначе используем Nmap для сетевого сканирования
        if not HAS_NMAP or not self.nm:
            print(f"{Fore.RED}[!] python-nmap не установлен. Используйте веб-сканирование через URL")
            return False
        
        profiles = {
            "quick": "-sV --top-ports 100 -T4 --min-rate 1000",
            "full": "-sV -sC -p- -O -T4",
            "vuln": "-sV -sC --script=vuln --script-args=unsafe=1 -p- -T4",
            "stealth": "-sS -Pn -T2 -f --max-retries 1 --min-rate 100",
            "web": "-sV -sC -p80,443,8080,8443 --script=http-*",
            "discovery": "-sV -sC -A -T4"
        }
        
        final_arguments = arguments if arguments else profiles.get(scan_profile, profiles["vuln"])
        
        try:
            print(f"{Fore.CYAN}[*] Запуск Nmap: {self.target}")
            print(f"{Fore.CYAN}[*] Профиль: {scan_profile}")
            
            self.nm.scan(self.target, arguments=final_arguments, sudo=False)
            print(f"{Fore.GREEN}[+] Сканирование завершено")
            return True
            
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка: {e}")
            return False
    
    def parse_results(self, search_cves: bool = True):
        """Парсинг результатов Nmap"""
        # Если это был веб-скан, результаты уже обработаны
        if self.is_url:
            return
        
        # Для Nmap сканирования
        if not self.nm:
            return
        
        try:
            ports_for_cve = []
            
            for host in self.nm.all_hosts():
                host_info = {
                    'host': host,
                    'status': self.nm[host].state(),
                    'hostname': self.nm[host].hostname() or '',
                    'ports': []
                }
                
                for proto in self.nm[host].all_protocols():
                    for port in self.nm[host][proto].keys():
                        port_data = self.nm[host][proto][port]
                        
                        port_info = {
                            'host': host,
                            'port': port,
                            'protocol': proto,
                            'state': port_data.get('state', 'unknown'),
                            'service': port_data.get('name', 'unknown'),
                            'product': port_data.get('product', ''),
                            'version': port_data.get('version', '')
                        }
                        
                        host_info['ports'].append(port_info)
                        
                        if port_info['state'] == 'open':
                            self.open_ports.append(port_info)
                            if search_cves and port_info['service'] != 'unknown':
                                ports_for_cve.append(port_info)
                
                self.discovered_hosts.append(host_info)
            
            if ports_for_cve:
                self._search_cves_parallel(ports_for_cve)
            
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка парсинга: {e}")
    
    def _search_cves_parallel(self, ports: List[Dict]):
        """Параллельный поиск CVE"""
        print(f"{Fore.CYAN}[*] Поиск CVE для {len(ports)} сервисов...")
        
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = {executor.submit(self._search_service_cves, p): p for p in ports}
            
            for future in as_completed(futures):
                try:
                    future.result()
                except Exception as e:
                    pass
    
    def _search_service_cves(self, port_info: Dict) -> List[CVEInfo]:
        """Поиск CVE для сервиса"""
        service = port_info.get('product') or port_info.get('service', '')
        version = port_info.get('version', '')
        
        if not service or service == 'unknown':
            return []
        
        cves = self.cve_searcher.search_cve(service, version if version else None)
        
        for cve in cves:
            self.vulnerabilities.append({
                'cve_id': cve.cve_id,
                'severity': cve.severity,
                'cvss_score': cve.cvss_score,
                'description': cve.description,
                'service': service,
                'version': version,
                'port': port_info['port'],
                'host': port_info['host']
            })
        
        return cves
    
    def get_summary(self) -> Dict:
        """Получить сводку"""
        severity_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0, 'Unknown': 0}
        
        all_vulns = self.vulnerabilities + self.web_vulnerabilities
        for vuln in all_vulns:
            severity = vuln.get('severity', 'Unknown')
            if severity in severity_counts:
                severity_counts[severity] += 1
        
        return {
            'target': self.target,
            'scan_duration': (datetime.now() - self.start_time).total_seconds(),
            'total_hosts': len(self.discovered_hosts),
            'total_ports': len(self.open_ports),
            'total_vulnerabilities': len(all_vulns),
            'severity_breakdown': severity_counts,
            'scan_profile': self.scan_profile,
            'is_web_scan': self.is_url
        }
    
    def get_json_report(self) -> Dict:
        """Получить JSON отчет"""
        summary = self.get_summary()
        
        if self.is_url:
            # Для веб-приложений
            return {
                'scan_info': {
                    'target': self.target,
                    'scan_id': self.scan_id,
                    'scan_datetime': self.start_time.isoformat(),
                    'scan_duration': summary['scan_duration'],
                    'profile': self.scan_profile,
                    'type': 'web_application'
                },
                'summary': summary,
                'web_info': self.web_info,
                'vulnerabilities': self.web_vulnerabilities
            }
        else:
            # Для сетевого сканирования
            return {
                'scan_info': {
                    'target': self.target,
                    'scan_id': self.scan_id,
                    'scan_datetime': self.start_time.isoformat(),
                    'scan_duration': summary['scan_duration'],
                    'profile': self.scan_profile,
                    'type': 'network_scan'
                },
                'summary': summary,
                'discovered_hosts': self.discovered_hosts,
                'open_ports': self.open_ports,
                'vulnerabilities': self.vulnerabilities
            }
    
    def _generate_txt_report(self) -> str:
        """Генерировать TXT отчет"""
        summary = self.get_summary()
        
        if self.is_url:
            # Отчет для веб-приложения
            lines = [
                "=" * 100,
                "ОТЧЕТ СКАНИРОВАНИЯ ВЕБ-ПРИЛОЖЕНИЯ".center(100),
                "=" * 100,
                "",
                f"URL: {self.target}",
                f"Host: {self.web_info.get('host', 'N/A')}",
                f"Port: {self.web_info.get('port', 'N/A')}",
                f"ID: {self.scan_id}",
                f"Дата: {self.start_time.strftime('%d.%m.%Y %H:%M:%S')}",
                f"Длительность: {summary['scan_duration']:.2f} сек",
                "",
                "ИНФОРМАЦИЯ О СЕРВЕРЕ",
                f"  Статус: HTTP {self.web_info.get('status_code', 'N/A')}",
                f"  Server: {self.web_info.get('server', 'Unknown')}",
                f"  Powered-By: {self.web_info.get('powered_by', 'N/A')}",
                f"  Технологии: {', '.join(self.web_info.get('technologies', ['N/A']))}",
                ""
            ]
            
            if self.web_info.get('ssl_expires'):
                lines.append(f"  SSL Expires: {self.web_info.get('ssl_expires')}")
            
            lines.extend([
                "",
                "СТАТИСТИКА",
                f"  Всего уязвимостей: {len(self.web_vulnerabilities)}",
                f"  Критические: {summary['severity_breakdown']['Critical']}",
                f"  Высокие: {summary['severity_breakdown']['High']}",
                f"  Средние: {summary['severity_breakdown']['Medium']}",
                f"  Низкие: {summary['severity_breakdown']['Low']}",
                ""
            ])
            
            if self.web_vulnerabilities:
                lines.append("НАЙДЕННЫЕ УЯЗВИМОСТИ")
                for i, vuln in enumerate(self.web_vulnerabilities, 1):
                    lines.append(f"\n[{i}] {vuln.get('type', 'Unknown')}")
                    lines.append(f"  Severity: {vuln.get('severity', 'Unknown')}")
                    lines.append(f"  Description: {vuln.get('description', 'N/A')}")
                    if vuln.get('header'):
                        lines.append(f"  Header: {vuln.get('header')}")
                    if vuln.get('url'):
                        lines.append(f"  URL: {vuln.get('url')}")
            
            lines.extend(["", "=" * 100])
            return "\n".join(lines)
        
        else:
            # Отчет для Nmap сканирования
            lines = [
                "=" * 100,
                "ОТЧЕТ СКАНИРОВАНИЯ NMAP".center(100),
                "=" * 100,
                "",
                f"Хост: {self.target}",
                f"ID: {self.scan_id}",
                f"Дата: {self.start_time.strftime('%d.%m.%Y %H:%M:%S')}",
                f"Профиль: {self.scan_profile}",
                f"Длительность: {summary['scan_duration']:.2f} сек",
                "",
                "СТАТИСТИКА",
                f"  Хостов: {summary['total_hosts']}",
                f"  Портов: {summary['total_ports']}",
                f"  Уязвимостей: {summary['total_vulnerabilities']}",
                "",
                "СЕРЬЕЗНОСТЬ",
                f"  Критические: {summary['severity_breakdown']['Critical']}",
                f"  Высокие: {summary['severity_breakdown']['High']}",
                f"  Средние: {summary['severity_breakdown']['Medium']}",
                f"  Низкие: {summary['severity_breakdown']['Low']}",
                ""
            ]
            
            if self.vulnerabilities:
                lines.append("УЯЗВИМОСТИ")
                for i, vuln in enumerate(self.vulnerabilities, 1):
                    lines.append(f"\n[{i}] {vuln.get('cve_id', 'N/A')}")
                    lines.append(f"  Severity: {vuln.get('severity', 'Unknown')}")
                    lines.append(f"  Service: {vuln.get('service', 'N/A')} {vuln.get('version', '')}")
                    lines.append(f"  Port: {vuln.get('port', 'N/A')}")
            
            lines.extend(["", "=" * 100])
            return "\n".join(lines)
    
    def _generate_word_report(self) -> Optional[str]:
        """Генерировать Word отчет"""
        if not HAS_DOCX:
            return None
        
        try:
            summary = self.get_summary()
            doc = Document()
            
            if self.is_url:
                # Отчет для веб-приложения
                doc.add_heading('ОТЧЕТ СКАНИРОВАНИЯ ВЕБ-ПРИЛОЖЕНИЯ', 0)
                
                doc.add_heading('ИНФОРМАЦИЯ', level=1)
                table = doc.add_table(rows=6, cols=2)
                table.style = 'Light Grid Accent 1'
                data = [
                    ('URL:', self.target),
                    ('Host:', self.web_info.get('host', 'N/A')),
                    ('Port:', str(self.web_info.get('port', 'N/A'))),
                    ('ID:', self.scan_id),
                    ('Дата:', self.start_time.strftime('%d.%m.%Y %H:%M:%S')),
                    ('Длительность:', f"{summary['scan_duration']:.2f} сек"),
                ]
                for i, (k, v) in enumerate(data):
                    table.rows[i].cells[0].text = k
                    table.rows[i].cells[1].text = str(v)
                
                doc.add_heading('СЕРВЕР', level=1)
                table = doc.add_table(rows=4, cols=2)
                table.style = 'Light Grid Accent 1'
                server_data = [
                    ('HTTP Статус:', str(self.web_info.get('status_code', 'N/A'))),
                    ('Server:', self.web_info.get('server', 'Unknown')),
                    ('Powered-By:', self.web_info.get('powered_by', 'N/A')),
                    ('Технологии:', ', '.join(self.web_info.get('technologies', ['N/A']))),
                ]
                for i, (k, v) in enumerate(server_data):
                    table.rows[i].cells[0].text = k
                    table.rows[i].cells[1].text = str(v)
                
                if self.web_vulnerabilities:
                    doc.add_heading('УЯЗВИМОСТИ', level=1)
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Light Grid Accent 1'
                    hdr = table.rows[0].cells
                    hdr[0].text = 'Тип'
                    hdr[1].text = 'Severity'
                    hdr[2].text = 'Описание'
                    hdr[3].text = 'URL/Заголовок'
                    
                    for vuln in self.web_vulnerabilities[:50]:
                        row = table.add_row().cells
                        row[0].text = vuln.get('type', 'N/A')
                        row[1].text = vuln.get('severity', 'Unknown')
                        row[2].text = vuln.get('description', 'N/A')[:50]
                        row[3].text = vuln.get('url', vuln.get('header', 'N/A'))
            
            else:
                # Отчет для Nmap
                doc.add_heading('ОТЧЕТ СКАНИРОВАНИЯ NMAP', 0)
                
                doc.add_heading('ИНФОРМАЦИЯ', level=1)
                table = doc.add_table(rows=5, cols=2)
                table.style = 'Light Grid Accent 1'
                data = [
                    ('Хост:', self.target),
                    ('ID:', self.scan_id),
                    ('Дата:', self.start_time.strftime('%d.%m.%Y %H:%M:%S')),
                    ('Профиль:', self.scan_profile),
                    ('Длительность:', f"{summary['scan_duration']:.2f} сек"),
                ]
                for i, (k, v) in enumerate(data):
                    table.rows[i].cells[0].text = k
                    table.rows[i].cells[1].text = str(v)
                
                doc.add_heading('УЯЗВИМОСТИ', level=1)
                if self.vulnerabilities:
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Light Grid Accent 1'
                    hdr = table.rows[0].cells
                    hdr[0].text = 'CVE'
                    hdr[1].text = 'Severity'
                    hdr[2].text = 'Service'
                    hdr[3].text = 'Port'
                    hdr[4].text = 'Version'
                    
                    for vuln in self.vulnerabilities[:50]:
                        row = table.add_row().cells
                        row[0].text = vuln.get('cve_id', 'N/A')
                        row[1].text = vuln.get('severity', 'Unknown')
                        row[2].text = vuln.get('service', 'N/A')
                        row[3].text = str(vuln.get('port', 'N/A'))
                        row[4].text = vuln.get('version', '')
            
            word_path = self.word_dir / f"{self.filename_base}.docx"
            doc.save(str(word_path))
            print(f"{Fore.GREEN}[+] Word отчет: {word_path}")
            return str(word_path)
        except Exception as e:
            print(f"{Fore.YELLOW}[!] Word ошибка: {e}")
            return None
    
    def print_report(self):
        """Вывести и сохранить отчеты"""
        txt_report = self._generate_txt_report()
        print("\n" + txt_report + "\n")
        
        self.save_json_report()
        self.save_txt_report()
        self._generate_word_report()
    
    def save_reports(self) -> Dict[str, str]:
        """Сохранить JSON и TXT отчеты (для совместимости с server.py)"""
        reports = {}
        
        # Сохраняем JSON
        json_path = self.save_json_report()
        if json_path:
            reports['json'] = json_path
        
        # Сохраняем TXT
        txt_path = self.save_txt_report()
        if txt_path:
            reports['txt'] = txt_path
        
        # Сохраняем WORD если доступен
        if HAS_DOCX:
            word_path = self._generate_word_report()
            if word_path:
                reports['word'] = word_path
        
        return reports

    def save_json_report(self) -> Optional[str]:
        """Сохранить JSON отчет"""
        try:
            json_path = self.json_dir / f"{self.filename_base}.json"
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(self.get_json_report(), f, indent=2, ensure_ascii=False)
            print(f"{Fore.GREEN}[+] JSON отчет сохранен: {json_path}")
            return str(json_path)
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка сохранения JSON: {e}")
            return None

    def save_txt_report(self) -> Optional[str]:
        """Сохранить TXT отчет"""
        try:
            txt_path = self.txt_dir / f"{self.filename_base}.txt"
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(self._generate_txt_report())
            print(f"{Fore.GREEN}[+] TXT отчет сохранен: {txt_path}")
            return str(txt_path)
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка сохранения TXT: {e}")
            return None

    def _count_by_severity(self) -> Dict[str, int]:
        """Подсчет уязвимостей по серьезности (для server.py)"""
        severity_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0, 'Unknown': 0}
        
        all_vulns = self.vulnerabilities + self.web_vulnerabilities
        for vuln in all_vulns:
            severity = vuln.get('severity', 'Unknown')
            if severity in severity_counts:
                severity_counts[severity] += 1
            else:
                severity_counts['Unknown'] += 1
        
        return severity_counts

    def _get_recommendations(self) -> List[str]:
        """Получить рекомендации (для server.py)"""
        if self.is_url:
            # Рекомендации для веб-приложений
            recommendations = [
                'Добавьте все рекомендуемые заголовки безопасности',
                'Включите HTTPS и HSTS',
                'Обновите все уязвимые компоненты',
                'Проведите аудит безопасности кода',
                'Используйте Web Application Firewall (WAF)'
            ]
            
            if self.web_vulnerabilities:
                critical_count = sum(1 for v in self.web_vulnerabilities if v.get('severity') == 'Critical')
                if critical_count > 0:
                    recommendations.insert(0, f'КРИТИЧНО: Закрытие обнаруженных {critical_count} критических уязвимостей!')
            
            return recommendations
        else:
            # Рекомендации для сетевого сканирования
            recommendations = [
                'Закройте ненужные открытые порты',
                'Обновите все обнаруженные сервисы',
                'Используйте минимально необходимые сервисы',
                'Настройте firewall для ограничения доступа',
                'Регулярно проводите сканирования портов'
            ]
            
            if self.vulnerabilities:
                recommendations.append('Обновите уязвимые сервисы как можно скорее')
            
            return recommendations

    def scan(self):
        """Метод для совместимости с callback"""
        return self.run_scan()

    def display_results(self):
        """Метод для совместимости с callback"""
        self.print_report()

    def _save_json_internal(self) -> Optional[str]:
        """Сохранить JSON отчет (внутренний метод)"""
        try:
            json_path = self.json_dir / f"{self.filename_base}.json"
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(self.get_json_report(), f, indent=2, ensure_ascii=False)
            print(f"{Fore.GREEN}[+] JSON отчет сохранен: {json_path}")
            return str(json_path)
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка сохранения JSON: {e}")
            return None

    def _save_txt_internal(self) -> Optional[str]:
        """Сохранить TXT отчет (внутренний метод)"""
        try:
            txt_path = self.txt_dir / f"{self.filename_base}.txt"
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(self._generate_txt_report())
            print(f"{Fore.GREEN}[+] TXT отчет сохранен: {txt_path}")
            return str(txt_path)
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка сохранения TXT: {e}")
            return None
    
    def save_json_report(self):
        """Сохранить JSON"""
        return self._save_json_internal()
    
    def save_txt_report(self):
        """Сохранить TXT"""
        return self._save_txt_internal()
    
    def save_reports(self) -> Dict[str, str]:
        """Сохранить JSON и TXT отчеты (для совместимости с server.py)"""
        reports = {}
        
        # Сохраняем JSON
        json_path = self._save_json_internal()
        if json_path:
            reports['json'] = json_path
        
        # Сохраняем TXT
        txt_path = self._save_txt_internal()
        if txt_path:
            reports['txt'] = txt_path
        
        # Сохраняем WORD если доступен
        if HAS_DOCX:
            word_path = self._generate_word_report()
            if word_path:
                reports['word'] = word_path
        
        return reports


def simple_scan(target: str, profile: str = "standard", reports_dir: str = None, nmap_profile: str = "standard") -> Dict:
    """Простая функция для server.py - поддержка URL и IP с профилями конфигурации
    
    Args:
        target: IP, домен или URL для сканирования
        profile: Профиль сканирования (deprecated)
        reports_dir: Директория для отчетов
        nmap_profile: Профиль из конфигурации (quick, standard, deep, aggressive)
    """
    scanner = NmapScanner(target=target, reports_dir=reports_dir)
    
    # Получаем аргументы из конфигурации профиля
    arguments = None
    if get_nmap_profile:
        try:
            cfg = get_nmap_profile(nmap_profile)
            arguments = cfg.get_full_arguments()
            print(f"[+] Используется профиль Nmap: {cfg.name}")
        except Exception as e:
            print(f"[!] Ошибка при загрузке профиля {nmap_profile}: {e}. Используем стандартные параметры.")
    
    if scanner.run_scan(scan_profile=profile, arguments=arguments):
        scanner.parse_results(search_cves=True)
        
        if scanner.is_url:
            # Для веб-приложений
            return {
                'status': 'success',
                'target': target,
                'scan_id': scanner.scan_id,
                'nmap_profile': nmap_profile,
                'summary': scanner.get_summary(),
                'web_info': scanner.web_info,
                'vulnerabilities': scanner.web_vulnerabilities,
                'type': 'web_application'
            }
        else:
            # Для сетевого сканирования
            return {
                'status': 'success',
                'target': target,
                'scan_id': scanner.scan_id,
                'nmap_profile': nmap_profile,
                'summary': scanner.get_summary(),
                'vulnerabilities': scanner.vulnerabilities,
                'open_ports': scanner.open_ports,
                'discovered_hosts': scanner.discovered_hosts,
                'type': 'network_scan'
            }
    else:
        return {'status': 'error', 'target': target, 'message': 'Scan failed'}


def main():
    """Основная функция"""
    parser = argparse.ArgumentParser(
        description='Advanced Nmap/Web Scanner with CVE API',
        epilog='Examples:\n  python nmap_scanner.py 192.168.1.1\n  python nmap_scanner.py http://127.0.0.1:5012/\n  python nmap_scanner.py example.com -p quick'
    )
    
    parser.add_argument('target', help='Target IP, domain or URL (http/https)')
    parser.add_argument('-p', '--profile', default='vuln', 
                       choices=['quick', 'full', 'vuln', 'stealth', 'web', 'discovery'],
                       help='Scan profile (for network scans)')
    parser.add_argument('-a', '--arguments', help='Custom Nmap arguments')
    parser.add_argument('-o', '--output', help='Output filename')
    parser.add_argument('-r', '--reports-dir', default='./reports', help='Reports directory')
    parser.add_argument('--nvd-api-key', help='NVD API key')
    
    args = parser.parse_args()
    
    print(f"{Fore.CYAN}╔{'═'*70}╗")
    print(f"║{'ADVANCED SCANNER (Network & Web) v4.0'.center(70)}║")
    print(f"╚{'═'*70}╝{Fore.RESET}\n")
    
    scanner = NmapScanner(
        target=args.target,
        output_base=args.output,
        reports_dir=args.reports_dir,
        nvd_api_key=args.nvd_api_key
    )
    
    if scanner.run_scan(arguments=args.arguments, scan_profile=args.profile):
        scanner.parse_results(search_cves=True)
        scanner.print_report()
        print(f"\n{Fore.GREEN}[✓] Успешно завершено!{Fore.RESET}\n")
        sys.exit(0)
    else:
        print(f"{Fore.RED}[!] Ошибка сканирования{Fore.RESET}\n")
        sys.exit(1)



if __name__ == "__main__":
    main()
